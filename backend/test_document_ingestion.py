"""Isolated extraction checks for TXT, selectable-text PDF, and DOCX uploads."""

from io import BytesIO

import fitz
from docx import Document

from document_ingestion import extract_document_text


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Unit 3: Database normalization")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Topic"
    table.rows[0].cells[1].text = "Normal forms"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _pdf_bytes() -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Question bank: explain ACID properties.")
    output = document.tobytes()
    document.close()
    return output


def _blank_pdf_bytes() -> bytes:
    document = fitz.open()
    document.new_page()
    output = document.tobytes()
    document.close()
    return output


def main() -> None:
    text, mime_type = extract_document_text("notice.txt", b"Mandatory lab at 10 AM", 5_000)
    assert text == "Mandatory lab at 10 AM"
    assert mime_type == "text/plain"

    docx_text, docx_mime = extract_document_text("unit-notes.docx", _docx_bytes(), 5_000)
    assert "Database normalization" in docx_text
    assert "Topic | Normal forms" in docx_text
    assert docx_mime.endswith("wordprocessingml.document")

    pdf_text, pdf_mime = extract_document_text("question-bank.pdf", _pdf_bytes(), 5_000)
    assert "ACID properties" in pdf_text
    assert pdf_mime == "application/pdf"

    try:
        extract_document_text("scan.pdf", _blank_pdf_bytes(), 5_000)
    except ValueError as error:
        assert "selectable text" in str(error)
    else:
        raise AssertionError("An image-only or blank PDF must be rejected.")

    try:
        extract_document_text("notes.pdf", _pdf_bytes(), 10)
    except ValueError as error:
        assert "more than" in str(error)
    else:
        raise AssertionError("Overlong extracted text must be rejected.")

    print("Document ingestion checks passed.")


if __name__ == "__main__":
    main()
