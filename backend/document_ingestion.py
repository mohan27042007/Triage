"""Bounded text extraction for student-uploaded TXT, PDF, and DOCX files."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile


SUPPORTED_DOCUMENT_TYPES = {".txt", ".pdf", ".docx"}
MAX_DOCUMENT_PAGES = 100
MAX_DOCX_UNCOMPRESSED_BYTES = 80 * 1024 * 1024

MIME_TYPES = {
    ".txt": "text/plain",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def extract_document_text(filename: str, file_bytes: bytes, max_text_chars: int) -> tuple[str, str]:
    """Return bounded selectable text and MIME metadata for one supported upload.

    This is intentionally extraction only: Triage does not execute document
    content, OCR image-only PDFs, or interpret macros. Callers retain the
    original bytes through the archive layer after extraction succeeds.
    """
    extension = Path(filename or "").suffix.lower()
    if extension not in SUPPORTED_DOCUMENT_TYPES:
        supported = ", ".join(sorted(SUPPORTED_DOCUMENT_TYPES))
        raise ValueError(f"The file must use one of: {supported}.")
    if not file_bytes:
        raise ValueError("The uploaded file is empty.")
    if max_text_chars < 1:
        raise ValueError("Document text limit must be positive.")

    if extension == ".txt":
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise ValueError("Text files must be UTF-8 encoded.") from exc
    elif extension == ".pdf":
        text = _extract_pdf_text(file_bytes)
    else:
        text = _extract_docx_text(file_bytes)

    normalized = "\n".join(line.rstrip() for line in text.splitlines()).strip()
    if not normalized:
        raise ValueError("Triage could not find selectable text in this document.")
    if len(normalized) > max_text_chars:
        raise ValueError(
            f"The document contains more than {max_text_chars:,} characters of extractable text. "
            "Upload a shorter excerpt or split the document."
        )
    return normalized, MIME_TYPES[extension]


def _extract_pdf_text(file_bytes: bytes) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("PDF uploads require the PyMuPDF dependency.") from exc
    try:
        document = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        raise ValueError("The PDF could not be opened.") from exc
    try:
        if document.needs_pass:
            raise ValueError("Password-protected PDFs are not supported.")
        if document.page_count > MAX_DOCUMENT_PAGES:
            raise ValueError(f"PDF uploads are limited to {MAX_DOCUMENT_PAGES} pages.")
        return "\n".join(page.get_text("text") for page in document)
    finally:
        document.close()


def _extract_docx_text(file_bytes: bytes) -> str:
    _validate_docx_archive(file_bytes)
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX uploads require the python-docx dependency.") from exc
    try:
        document = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise ValueError("The DOCX file could not be opened.") from exc

    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _validate_docx_archive(file_bytes: bytes) -> None:
    try:
        with ZipFile(BytesIO(file_bytes)) as archive:
            uncompressed_bytes = sum(entry.file_size for entry in archive.infolist())
            if uncompressed_bytes > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ValueError("The DOCX expands beyond Triage's safe extraction limit.")
            if "word/document.xml" not in archive.namelist():
                raise ValueError("The file is not a valid DOCX document.")
    except BadZipFile as exc:
        raise ValueError("The DOCX file is invalid or corrupted.") from exc
